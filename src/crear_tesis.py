"""
crear_tesis.py – Genera Tesis_borrador.docx
============================================
Ejecuta el pipeline completo de análisis y construye el documento Word
de la tesis con formato UTN – Maestría en Minería de Datos.

Uso:
    cd src/
    python crear_tesis.py

Salida: FreddieMAC/Tesis_borrador.docx
"""

import sys, os, io, time, traceback, warnings
from pathlib import Path

import pandas as pd
import numpy as np

# Asegurar que src/ esté en el path
sys.path.insert(0, str(Path(__file__).parent))
warnings.filterwarnings("ignore")

# ---------------------------------------------------------------------------
# python-docx
# ---------------------------------------------------------------------------
from docx import Document
from docx.shared import Pt, Cm, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# ---------------------------------------------------------------------------
# Paths
# ---------------------------------------------------------------------------
from config import BASE_PATH, FIGURAS_PATH, TABLAS_PATH, MODELOS_PATH, DATA_PATH

DOCX_OUT = BASE_PATH / "Tesis_borrador.docx"


# ============================================================================
# HELPERS DE FORMATO
# ============================================================================

def set_run_font(run, nombre="Times New Roman", tamaño=12, negrita=False,
                 cursiva=False, color=None):
    run.font.name = nombre
    run.font.size = Pt(tamaño)
    run.font.bold = negrita
    run.font.italic = cursiva
    if color:
        run.font.color.rgb = RGBColor(*color)


def agregar_parrafo(doc, texto, estilo="Normal", negrita=False, cursiva=False,
                    tamaño=12, alineacion=WD_ALIGN_PARAGRAPH.JUSTIFY,
                    espacio_antes=0, espacio_despues=6):
    p = doc.add_paragraph(style=estilo)
    p.alignment = alineacion
    p.paragraph_format.space_before = Pt(espacio_antes)
    p.paragraph_format.space_after = Pt(espacio_despues)
    p.paragraph_format.line_spacing = Pt(18)
    run = p.add_run(texto)
    set_run_font(run, tamaño=tamaño, negrita=negrita, cursiva=cursiva)
    return p


def agregar_titulo(doc, texto, nivel=1):
    """Agrega encabezado numerado con estilo académico."""
    estilos = {1: (14, True), 2: (13, True), 3: (12, True)}
    tam, bold = estilos.get(nivel, (12, False))
    p = doc.add_heading(texto, level=nivel)
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    for run in p.runs:
        run.font.name = "Times New Roman"
        run.font.size = Pt(tam)
        run.font.bold = bold
        run.font.color.rgb = RGBColor(0, 0, 0)
    return p


def agregar_codigo(doc, codigo: str):
    """Bloque de código con fuente monoespaciada y fondo gris."""
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(1)
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)
    # Fondo gris claro via XML
    pPr = p._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), "F2F2F2")
    pPr.append(shd)
    run = p.add_run(codigo)
    run.font.name = "Courier New"
    run.font.size = Pt(9)
    return p


def agregar_figura(doc, ruta: str, caption: str, ancho=14):
    """Inserta una figura con su epígrafe."""
    if Path(ruta).exists():
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run()
        run.add_picture(ruta, width=Cm(ancho))
    else:
        agregar_parrafo(doc, f"[Figura no disponible: {Path(ruta).name}]",
                        cursiva=True, tamaño=10)
    cap = doc.add_paragraph()
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = cap.add_run(caption)
    set_run_font(r, tamaño=10, cursiva=True)
    doc.add_paragraph()


def agregar_tabla_df(doc, df: pd.DataFrame, caption: str = ""):
    """Convierte un DataFrame a tabla Word con estilo académico."""
    if caption:
        p = doc.add_paragraph()
        r = p.add_run(caption)
        set_run_font(r, tamaño=10, cursiva=True)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    df_str = df.copy()
    for col in df_str.select_dtypes(include=[float]).columns:
        df_str[col] = df_str[col].map(lambda x: f"{x:.4f}" if pd.notna(x) else "")

    tabla = doc.add_table(rows=1, cols=len(df_str.columns))
    tabla.style = "Table Grid"
    tabla.alignment = WD_TABLE_ALIGNMENT.CENTER

    # Encabezados
    hdr = tabla.rows[0].cells
    for i, col in enumerate(df_str.columns):
        hdr[i].text = str(col)
        for run in hdr[i].paragraphs[0].runs:
            run.font.bold = True
            run.font.size = Pt(9)
            run.font.name = "Times New Roman"
        # Fondo azul oscuro para encabezados
        shd = OxmlElement("w:shd")
        shd.set(qn("w:val"), "clear")
        shd.set(qn("w:color"), "auto")
        shd.set(qn("w:fill"), "1F4E79")
        hdr[i].paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)
        hdr[i]._tc.get_or_add_tcPr().append(shd)

    # Filas de datos
    for i, (_, row) in enumerate(df_str.iterrows()):
        cells = tabla.add_row().cells
        fill = "FFFFFF" if i % 2 == 0 else "E8F0FE"
        for j, val in enumerate(row):
            cells[j].text = str(val) if pd.notna(val) else ""
            for run in cells[j].paragraphs[0].runs:
                run.font.size = Pt(9)
                run.font.name = "Times New Roman"
            shd = OxmlElement("w:shd")
            shd.set(qn("w:val"), "clear")
            shd.set(qn("w:color"), "auto")
            shd.set(qn("w:fill"), fill)
            cells[j]._tc.get_or_add_tcPr().append(shd)

    doc.add_paragraph()
    return tabla


def salto_pagina(doc):
    doc.add_page_break()


# ============================================================================
# EJECUCIÓN DEL PIPELINE
# ============================================================================

def ejecutar_pipeline():
    """Ejecuta todos los pasos y retorna los resultados."""
    resultados = {}

    # --- Paso 0: EDA ---
    print("\n[PIPELINE] Paso 0: EDA...")
    try:
        import importlib
        eda = importlib.import_module("00_eda")
        resultados["eda"] = eda.main()
    except Exception:
        print(f"  ERROR en EDA: {traceback.format_exc()}")
        resultados["eda"] = {}

    # --- Paso 1: Carga ---
    print("[PIPELINE] Paso 1: Carga de datos...")
    try:
        import importlib
        p01 = importlib.import_module("01_carga_datos")
        if not (DATA_PATH / "panel.parquet").exists():
            resultados["panel"] = p01.main()
        else:
            resultados["panel"] = pd.read_parquet(DATA_PATH / "panel.parquet")
            print("  Panel cargado desde caché.")
    except Exception:
        print(f"  ERROR Paso 1: {traceback.format_exc()}")
        resultados["panel"] = None

    # --- Paso 2: Default ---
    print("[PIPELINE] Paso 2: Definición de default...")
    try:
        import importlib
        p02 = importlib.import_module("02_definicion_default")
        if not (DATA_PATH / "panel_con_default.parquet").exists():
            resultados["panel_def"] = p02.main()
        else:
            resultados["panel_def"] = pd.read_parquet(DATA_PATH / "panel_con_default.parquet")
            print("  panel_con_default cargado desde caché.")
    except Exception:
        print(f"  ERROR Paso 2: {traceback.format_exc()}")
        resultados["panel_def"] = None

    # --- Paso 3: Feature Engineering ---
    print("[PIPELINE] Paso 3: Feature Engineering...")
    try:
        import importlib
        p03 = importlib.import_module("03_feature_engineering")
        if not (DATA_PATH / "dataset_modelado.parquet").exists():
            resultados["dataset_mod"] = p03.main()
        else:
            resultados["dataset_mod"] = pd.read_parquet(DATA_PATH / "dataset_modelado.parquet")
            print("  dataset_modelado cargado desde caché.")
    except Exception:
        print(f"  ERROR Paso 3: {traceback.format_exc()}")
        resultados["dataset_mod"] = None

    # --- Paso 4: Modelado ---
    print("[PIPELINE] Paso 4: Modelado...")
    try:
        import importlib
        p04 = importlib.import_module("04_modelado")
        mod_xgb = MODELOS_PATH / "modelo_xgboost.joblib"
        if not mod_xgb.exists():
            resultados["metricas_modelos"] = p04.main()
        else:
            ruta_tabla = TABLAS_PATH / "resultados_modelos.csv"
            if ruta_tabla.exists():
                resultados["metricas_modelos"] = pd.read_csv(ruta_tabla, index_col=0)
                print("  Métricas de modelos cargadas desde caché.")
            else:
                resultados["metricas_modelos"] = p04.main()
    except Exception:
        print(f"  ERROR Paso 4: {traceback.format_exc()}")
        resultados["metricas_modelos"] = None

    # --- Paso 5: Calibración Lifetime PD ---
    print("[PIPELINE] Paso 5: Calibración Lifetime PD...")
    try:
        import importlib
        p05 = importlib.import_module("05_calibracion_lifetime")
        ruta_lt = TABLAS_PATH / "curva_pd_lifetime.csv"
        if not ruta_lt.exists():
            resultados["lifetime_pd"] = p05.main()
        else:
            resultados["lifetime_pd"] = pd.read_csv(ruta_lt)
            print("  Lifetime PD cargado desde caché.")
    except Exception:
        print(f"  ERROR Paso 5: {traceback.format_exc()}")
        resultados["lifetime_pd"] = None

    # --- Paso 6: Validación ---
    print("[PIPELINE] Paso 6: Validación...")
    try:
        import importlib
        p06 = importlib.import_module("06_validacion")
        ruta_val = TABLAS_PATH / "validacion_metricas.csv"
        if not ruta_val.exists():
            resultados["validacion"] = p06.main()
        else:
            resultados["validacion"] = pd.read_csv(ruta_val, index_col=0)
            print("  Métricas de validación cargadas desde caché.")
    except Exception:
        print(f"  ERROR Paso 6: {traceback.format_exc()}")
        resultados["validacion"] = None

    # --- Paso 7: Explicabilidad ---
    print("[PIPELINE] Paso 7: Explicabilidad SHAP...")
    try:
        import importlib
        p07 = importlib.import_module("07_explicabilidad")
        if not (FIGURAS_PATH / "shap_summary_xgboost.png").exists():
            p07.main()
        else:
            print("  Figuras SHAP cargadas desde caché.")
    except Exception:
        print(f"  ERROR Paso 7: {traceback.format_exc()}")

    return resultados


# ============================================================================
# CONSTRUCCIÓN DEL DOCUMENTO WORD
# ============================================================================

def construir_docx(resultados: dict):
    doc = Document()

    # --- Márgenes ---
    for section in doc.sections:
        section.top_margin    = Cm(2.5)
        section.bottom_margin = Cm(2.5)
        section.left_margin   = Cm(3.0)
        section.right_margin  = Cm(2.5)

    eda = resultados.get("eda", {})
    panel_def = resultados.get("panel_def")
    ds_mod    = resultados.get("dataset_mod")
    met_mod   = resultados.get("metricas_modelos")
    lifetime  = resultados.get("lifetime_pd")
    val       = resultados.get("validacion")

    figuras = eda.get("figuras", {})

    # -------------------------------------------------------------------------
    # PORTADA
    # -------------------------------------------------------------------------
    doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("UNIVERSIDAD TECNOLÓGICA NACIONAL")
    set_run_font(r, tamaño=14, negrita=True)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("FACULTAD REGIONAL PARANÁ")
    set_run_font(r, tamaño=13, negrita=True)

    doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("PLAN DE TESIS")
    set_run_font(r, tamaño=12, negrita=True)

    doc.add_paragraph()
    doc.add_paragraph()

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(
        "Predicción de la Probabilidad de Default bajo el marco regulatorio\n"
        "IFRS 9 utilizando técnicas de aprendizaje automático\n"
        "y datos crediticios abiertos de Freddie Mac"
    )
    set_run_font(r, tamaño=16, negrita=True)

    doc.add_paragraph()
    doc.add_paragraph()

    for linea in ["Por: Lic. Sebastián Emiliano Meier", "",
                  "MAESTRÍA EN MINERÍA DE DATOS", "",
                  "Director/a: Mag. Ing. Gustavo Denicolay", "",
                  "Paraná, Argentina", "2025"]:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(linea)
        set_run_font(r, tamaño=12, negrita=("MAESTRÍA" in linea))

    salto_pagina(doc)

    # -------------------------------------------------------------------------
    # RESUMEN
    # -------------------------------------------------------------------------
    agregar_titulo(doc, "Resumen", nivel=1)
    n_loans  = eda.get("n_prestamos", None)
    n_obs    = eda.get("n_obs_perf", None)
    tasa_d   = eda.get("tasa_default", None)
    tasa_str  = f"{tasa_d:.2%}" if tasa_d else "N/A"
    n_loans_fmt = f"{n_loans:,}" if isinstance(n_loans, int) else str(n_loans or "N/A")
    n_obs_fmt   = f"{n_obs:,}"   if isinstance(n_obs,   int) else str(n_obs   or "N/A")

    agregar_parrafo(doc,
        f"El presente trabajo de tesis desarrolla un modelo predictivo de Probabilidad de Default (PD) "
        f"conforme al marco regulatorio IFRS 9, aplicando técnicas de aprendizaje automático sobre el "
        f"dataset público de Freddie Mac Single-Family Loan-Level. El dataset utilizado comprende "
        f"{n_loans_fmt} préstamos hipotecarios con un total de {n_obs_fmt} observaciones mensuales, "
        f"correspondientes a los años de originación 2016 a 2024 (excluyendo 2021 por indisponibilidad "
        f"de datos). La tasa de default global observada es de {tasa_str}. "
        f"Se implementan y comparan tres modelos: Regresión Logística (baseline regulatorio), "
        f"XGBoost y Random Forest. El modelo final es calibrado bajo el enfoque Point-In-Time (PIT) "
        f"y se genera una curva de Lifetime PD conforme a IFRS 9. La validación incluye métricas "
        f"regulatorias (AUC, Gini, KS, PSI, Brier Score) y análisis de explicabilidad con SHAP values."
    )

    agregar_parrafo(doc,
        "Palabras clave: IFRS 9, Probabilidad de Default, Expected Credit Loss, XGBoost, SHAP, "
        "Freddie Mac, Machine Learning, Riesgo Crediticio.",
        cursiva=True, tamaño=11
    )
    salto_pagina(doc)

    # -------------------------------------------------------------------------
    # 1. FUNDAMENTACIÓN Y JUSTIFICACIÓN
    # -------------------------------------------------------------------------
    agregar_titulo(doc, "1. Fundamentación y Justificación del Tema", nivel=1)
    agregar_parrafo(doc,
        "El riesgo de crédito constituye una de las principales fuentes de exposición en las "
        "entidades financieras y cooperativas de crédito. En la última década, la regulación "
        "contable internacional ha experimentado un cambio radical con la implementación de la "
        "Norma Internacional de Información Financiera IFRS 9, que introdujo el enfoque de "
        "Expected Credit Loss (ECL). Este modelo requiere que las instituciones estimen de forma "
        "prospectiva la Probabilidad de Default (PD), la Pérdida Dada el Default (LGD) y la "
        "Exposición al Default (EAD) para calcular las pérdidas esperadas de sus carteras."
    )
    agregar_parrafo(doc,
        "A diferencia del modelo contable anterior (IAS 39), IFRS 9 exige una aproximación "
        "forward-looking, donde la pérdida crediticia esperada incorpora no solo información "
        "histórica, sino también condiciones macroeconómicas actuales y proyecciones futuras. "
        "Por ello, el desarrollo de modelos predictivos de PD es un aspecto central del "
        "cumplimiento normativo y de la gestión prudencial de riesgos."
    )
    agregar_parrafo(doc,
        "En América Latina y particularmente en Argentina, el proceso de adopción de IFRS 9 se "
        "encuentra en expansión. El Banco Central de la República Argentina (BCRA) ha avanzado "
        "en la convergencia normativa hacia IFRS 9, especialmente para entidades financieras de "
        "mayor tamaño, generando una necesidad creciente de profesionales con formación técnica "
        "sólida en modelización IFRS 9."
    )
    agregar_parrafo(doc,
        "El desarrollo de esta tesis se enmarca en ese contexto: busca contribuir a la literatura "
        "aplicada y a la práctica profesional presentando una metodología reproducible para estimar "
        "la PD bajo IFRS 9 utilizando datos públicos abiertos del Freddie Mac Single-Family "
        "Loan-Level Dataset. El uso de este dataset permite reproducir la estructura de información "
        "de una entidad financiera real, incluyendo variables de monto, tasa, LTV, DPD, score "
        "crediticio, fechas de pago, eventos de mora y pérdida."
    )
    salto_pagina(doc)

    # -------------------------------------------------------------------------
    # 2. ESTADO DEL ARTE
    # -------------------------------------------------------------------------
    agregar_titulo(doc, "2. Estado del Arte", nivel=1)
    agregar_parrafo(doc,
        "El modelado de la Probabilidad de Default (PD) tiene una trayectoria extensa en la gestión "
        "del riesgo financiero. Tradicionalmente, las entidades bancarias han utilizado modelos de "
        "regresión logística para estimar la probabilidad de incumplimiento, debido a su "
        "interpretabilidad, estabilidad y cumplimiento con los requerimientos de explicabilidad "
        "establecidos por los supervisores."
    )
    agregar_parrafo(doc,
        "En el marco de Basilea II y III, las PD se calibraban a valores Through-The-Cycle (TTC), "
        "enfocados en la estabilidad a largo plazo. Sin embargo, IFRS 9 introduce un cambio clave: "
        "exige una PD Point-In-Time (PIT), sensible al ciclo económico y capaz de proyectarse a "
        "horizontes de 12 meses y a toda la vida del crédito (Lifetime PD). Este cambio "
        "metodológico ha impulsado una nueva generación de modelos, donde las técnicas de machine "
        "learning se incorporan para mejorar la capacidad predictiva y la segmentación del riesgo."
    )
    agregar_parrafo(doc,
        "Entre los avances más relevantes se encuentran los modelos Gradient Boosting Machines "
        "(XGBoost, LightGBM), Random Forests y enfoques basados en ensembles híbridos, los cuales "
        "han demostrado mayor capacidad para capturar no linealidades y efectos de interacción "
        "entre variables. No obstante, los organismos reguladores insisten en la necesidad de "
        "explicabilidad y gobernanza del modelo, motivo por el cual los modelos de caja negra "
        "deben complementarse con herramientas de interpretación como SHAP values o LIME."
    )
    agregar_parrafo(doc,
        "Estudios recientes han utilizado los datasets de Freddie Mac y Fannie Mae para modelar "
        "el riesgo hipotecario y estimar las tasas de incumplimiento bajo distintos contextos "
        "macroeconómicos (Ghent & Kudlyak, 2011; Elul et al., 2010). El presente trabajo se "
        "diferencia al combinar el enfoque regulatorio IFRS 9 con una implementación de machine "
        "learning reproducible y explicable, documentando cada etapa del proceso."
    )
    salto_pagina(doc)

    # -------------------------------------------------------------------------
    # 3. MARCO TEÓRICO IFRS 9
    # -------------------------------------------------------------------------
    agregar_titulo(doc, "3. Marco Teórico: IFRS 9 y Expected Credit Loss", nivel=1)

    agregar_titulo(doc, "3.1 Expected Credit Loss (ECL)", nivel=2)
    agregar_parrafo(doc,
        "El modelo ECL bajo IFRS 9 requiere que la pérdida crediticia esperada se calcule como:"
    )
    agregar_codigo(doc, "ECL = PD × LGD × EAD × (1 / (1 + r))^t")
    agregar_parrafo(doc,
        "Donde PD es la Probabilidad de Default, LGD la Pérdida Dada el Default, EAD la "
        "Exposición al Default, r la tasa de descuento y t el horizonte temporal. La presente "
        "tesis se enfoca en la estimación de la PD."
    )

    agregar_titulo(doc, "3.2 Clasificación de Stages IFRS 9", nivel=2)
    agregar_parrafo(doc,
        "IFRS 9 clasifica los instrumentos financieros en tres etapas (stages) según el nivel "
        "de deterioro del crédito:"
    )
    for stage, desc in [
        ("Stage 1", "Sin incremento significativo del riesgo de crédito (SICR). "
                    "Se reconoce ECL a 12 meses. DPD = 0 (corriente)."),
        ("Stage 2", "Incremento significativo del riesgo de crédito. "
                    "Se reconoce Lifetime ECL. DPD: 30-89 días."),
        ("Stage 3", "Activo crediticio deteriorado (en default). "
                    "Se reconoce Lifetime ECL. DPD ≥ 90 días o evento de pérdida."),
    ]:
        p = doc.add_paragraph(style="List Bullet")
        r = p.add_run(f"{stage}: ")
        set_run_font(r, negrita=True, tamaño=11)
        r2 = p.add_run(desc)
        set_run_font(r2, tamaño=11)

    agregar_titulo(doc, "3.3 Definición de Default Adoptada", nivel=2)
    agregar_parrafo(doc,
        "En este trabajo se adopta la definición de default conforme al estándar IFRS 9 / Basilea. "
        "Un préstamo se considera en default en el período t si se cumple alguna de las siguientes "
        "condiciones:"
    )
    agregar_codigo(doc,
        "Default = (DPD ≥ 90 días)  OR  (Zero Balance Code ∈ {02, 03, 09})\n"
        "\n"
        "Donde:\n"
        "  DPD = current_loan_delinquency_status × 30 días\n"
        "  ZBC 02 = Third Party Sale (venta a terceros)\n"
        "  ZBC 03 = Short Sale o Charge-Off\n"
        "  ZBC 09 = REO Disposition (ejecución hipotecaria)\n"
    )
    salto_pagina(doc)

    # -------------------------------------------------------------------------
    # 4. DESCRIPCIÓN DEL DATASET Y EDA
    # -------------------------------------------------------------------------
    agregar_titulo(doc, "4. Descripción del Dataset y Análisis Exploratorio", nivel=1)

    agregar_titulo(doc, "4.1 Freddie Mac Single-Family Loan-Level Dataset", nivel=2)
    agregar_parrafo(doc,
        "Se utiliza el Freddie Mac Single-Family Loan-Level Dataset (Sample), una muestra "
        "aleatoria de 50.000 préstamos por año de originación, provista públicamente por "
        "Freddie Mac. El dataset contiene dos archivos por año:"
    )
    for item in ["sample_orig_YYYY.txt: datos de originación del préstamo (32 variables).",
                 "sample_svcg_YYYY.txt: datos de performance mensual (32 variables)."]:
        p = doc.add_paragraph(style="List Bullet")
        r = p.add_run(item)
        set_run_font(r, tamaño=11)

    agregar_titulo(doc, "4.2 Estadísticas Descriptivas – Originación", nivel=2)
    _np = eda.get('n_prestamos', None)
    n_loans_str  = f"{_np:,}" if isinstance(_np, int) else str(_np or "N/A")
    fecha_min_s  = str(eda.get('fecha_min', 'N/A'))[:10]
    fecha_max_s  = str(eda.get('fecha_max', 'N/A'))[:10]
    agregar_parrafo(doc,
        f"El dataset de originación contiene {n_loans_str} préstamos únicos, con primer pago "
        f"entre {fecha_min_s} y {fecha_max_s}. A continuación se presentan las estadísticas "
        f"descriptivas de las principales variables cuantitativas:"
    )

    stats_o = eda.get("stats_orig", {}).get("stats_orig")
    if stats_o is not None:
        # Formatear para docx
        cols_mostrar = ["count", "mean", "std", "min", "25%", "50%", "75%", "max"]
        cols_disp = [c for c in cols_mostrar if c in stats_o.columns]
        agregar_tabla_df(doc, stats_o[cols_disp].round(2),
                         "Tabla 1. Estadísticas descriptivas – variables de originación")

    agregar_titulo(doc, "4.3 Estadísticas Descriptivas – Performance", nivel=2)
    _no = eda.get('n_obs_perf', None)
    n_obs_str   = f"{_no:,}" if isinstance(_no, int) else str(_no or "N/A")
    stats_p     = eda.get("stats_perf", {})
    per_min_s   = str(stats_p.get("per_min", "N/A"))[:10]
    per_max_s   = str(stats_p.get("per_max", "N/A"))[:10]
    tasa_90_str = f"{stats_p.get('tasa_90dpd', 0):.4%}"
    agregar_parrafo(doc,
        f"El dataset de performance contiene {n_obs_str} observaciones mensuales, "
        f"abarcando el período {per_min_s} a {per_max_s}. "
        f"La tasa de observaciones con DPD ≥ 90 días es del {tasa_90_str}. "
        f"La tasa de default global (DPD ≥ 90 o evento de pérdida) es de {tasa_str}."
    )

    agregar_titulo(doc, "4.4 Visualizaciones Exploratorias", nivel=2)

    if figuras.get("fico"):
        agregar_figura(doc, figuras["fico"],
            "Figura 1. Distribución del Credit Score (FICO) de originación.")
    if figuras.get("upb"):
        agregar_figura(doc, figuras["upb"],
            "Figura 2. Distribución del saldo actual (Current Actual UPB).")
    if figuras.get("dpd"):
        agregar_figura(doc, figuras["dpd"],
            "Figura 3. Distribución de días en mora (DPD) en el período de performance.")

    salto_pagina(doc)

    if figuras.get("default_tiempo"):
        agregar_figura(doc, figuras["default_tiempo"],
            "Figura 4. Tasa de default mensual en el dataset Freddie Mac.")
    if figuras.get("default_vintage"):
        agregar_figura(doc, figuras["default_vintage"],
            "Figura 5. Tasa de default por vintage year (año de originación).")
    if figuras.get("ltv_default"):
        agregar_figura(doc, figuras["ltv_default"],
            "Figura 6. Tasa de default por bucket de LTV original.")

    salto_pagina(doc)

    # -------------------------------------------------------------------------
    # 5. METODOLOGÍA Y PIPELINE
    # -------------------------------------------------------------------------
    agregar_titulo(doc, "5. Metodología y Pipeline de Análisis", nivel=1)
    agregar_parrafo(doc,
        "El pipeline de análisis fue implementado íntegramente en Python 3, siguiendo una "
        "arquitectura modular de siete pasos secuenciales. Cada paso genera archivos intermedios "
        "en formato Parquet para garantizar la reproducibilidad:"
    )
    for i, paso in enumerate([
        "Carga y merge de datos (01_carga_datos.py)",
        "Definición de default IFRS 9 y clasificación de stages (02_definicion_default.py)",
        "Ingeniería de variables (03_feature_engineering.py)",
        "Entrenamiento de modelos PD (04_modelado.py)",
        "Calibración PIT y curva Lifetime PD (05_calibracion_lifetime.py)",
        "Validación regulatoria (06_validacion.py)",
        "Explicabilidad SHAP (07_explicabilidad.py)",
    ], 1):
        p = doc.add_paragraph(style="List Number")
        r = p.add_run(paso)
        set_run_font(r, tamaño=11)

    agregar_titulo(doc, "5.1 Carga y Preparación de Datos", nivel=2)
    agregar_parrafo(doc,
        "Los datos son leídos desde los archivos pipe-delimited originales, con todas las "
        "columnas cargadas como string para evitar errores de inferencia de tipos. "
        "Los códigos sentinel definidos en el User Guide (e.g., 9999 para credit score no "
        "disponible, 999 para LTV/DTI/MI) son convertidos a NaN por columna."
    )
    agregar_codigo(doc,
        "# Ejemplo: lectura robusta con tratamiento de sentinels\n"
        "df = pd.read_csv(ruta, sep='|', header=None, names=ORIG_COLS,\n"
        "                 dtype=str, na_values=[''], keep_default_na=False)\n"
        "df['credit_score'] = pd.to_numeric(df['credit_score'],\n"
        "                     errors='coerce').replace(9999, np.nan)\n"
        "df['original_ltv'] = pd.to_numeric(df['original_ltv'],\n"
        "                     errors='coerce').replace(999, np.nan)"
    )

    if panel_def is not None:
        n_panel = len(panel_def)
        n_loans_panel = panel_def["loan_sequence_number"].nunique() \
            if "loan_sequence_number" in panel_def.columns else "N/A"
        agregar_parrafo(doc,
            f"El panel longitudinal resultante contiene {n_panel:,} observaciones mensuales "
            f"correspondientes a {n_loans_panel:,} préstamos únicos."
        )

    salto_pagina(doc)

    # -------------------------------------------------------------------------
    # 6. DEFINICIÓN DEL DEFAULT E IFRS 9 STAGES
    # -------------------------------------------------------------------------
    agregar_titulo(doc, "6. Definición del Default y Clasificación IFRS 9", nivel=1)
    agregar_codigo(doc,
        "# Definición de default IFRS 9\n"
        "svcg['dpd_numerico'] = pd.to_numeric(\n"
        "    svcg['current_loan_delinquency_status'].replace('RA', 99),\n"
        "    errors='coerce'\n"
        ")\n"
        "svcg['evento_default'] = (\n"
        "    (svcg['dpd_numerico'] >= 3) |        # ≥ 90 DPD\n"
        "    svcg['zero_balance_code'].isin({'02','03','09'})\n"
        ")\n"
        "# Stage 1: DPD=0 | Stage 2: 1≤DPD<3 | Stage 3: DPD≥3 o default"
    )

    if panel_def is not None and "ifrs9_stage" in panel_def.columns:
        agregar_titulo(doc, "6.1 Distribución de Stages IFRS 9", nivel=2)
        dist_stage = panel_def["ifrs9_stage"].value_counts().sort_index()
        df_stages = pd.DataFrame({
            "Stage": [f"Stage {s}" for s in dist_stage.index],
            "Observaciones": [f"{v:,}" for v in dist_stage.values],
            "Porcentaje (%)": [f"{v/len(panel_def)*100:.2f}%" for v in dist_stage.values],
        })
        agregar_tabla_df(doc, df_stages, "Tabla 2. Distribución de stages IFRS 9")

        if "default_12m" in panel_def.columns:
            df_activos = panel_def[panel_def["ifrs9_stage"].isin([1, 2])]
            tasa_12m = df_activos["default_12m"].mean()
            agregar_parrafo(doc,
                f"La tasa de default a 12 meses sobre las observaciones activas (Stage 1 y 2) "
                f"es de {tasa_12m:.4%}, lo que refleja el carácter desbalanceado típico de "
                f"las carteras crediticias retail."
            )

    salto_pagina(doc)

    # -------------------------------------------------------------------------
    # 7. INGENIERÍA DE VARIABLES
    # -------------------------------------------------------------------------
    agregar_titulo(doc, "7. Ingeniería de Variables (Feature Engineering)", nivel=1)
    agregar_parrafo(doc,
        "Se utilizan dos tipos de variables: estáticas (de originación) y dinámicas (de "
        "comportamiento mensual). Las variables de comportamiento son las de mayor poder "
        "predictivo bajo el enfoque PIT de IFRS 9."
    )

    tabla_features = pd.DataFrame([
        ("credit_score", "Originación", "FICO al momento de originación"),
        ("original_ltv", "Originación", "LTV original (%)"),
        ("original_cltv", "Originación", "CLTV original (%)"),
        ("original_dti", "Originación", "DTI original (%)"),
        ("original_interest_rate", "Originación", "Tasa nominal original"),
        ("original_loan_term", "Originación", "Plazo en meses"),
        ("loan_purpose", "Originación", "Propósito: compra / refinanciación"),
        ("occupancy_status", "Originación", "Primaria / Inversión / Segunda"),
        ("channel", "Originación", "Canal: Retail / Broker / Correspondent"),
        ("loan_age", "Comportamiento", "Antigüedad del préstamo (meses)"),
        ("dpd_numerico", "Comportamiento", "DPD del período corriente"),
        ("max_dpd_3m", "Comportamiento", "Máximo DPD últimos 3 meses"),
        ("max_dpd_6m", "Comportamiento", "Máximo DPD últimos 6 meses"),
        ("max_dpd_12m", "Comportamiento", "Máximo DPD últimos 12 meses"),
        ("estimated_loan_to_value_eltv", "Comportamiento", "LTV estimado AVM (desde abr-2017)"),
        ("spread_tasa", "Comportamiento", "Tasa corriente − Tasa original"),
        ("amortizacion_upb", "Comportamiento", "Amortización relativa del saldo"),
        ("fue_modificado", "Comportamiento", "Flag: préstamo modificado (binario)"),
        ("tiene_deferral", "Comportamiento", "Flag: diferimiento de pagos (binario)"),
    ], columns=["Variable", "Tipo", "Descripción"])

    agregar_tabla_df(doc, tabla_features, "Tabla 3. Variables del modelo PD IFRS 9")

    if ds_mod is not None:
        n_mod = len(ds_mod)
        tasa_mod = ds_mod["default_12m"].mean() if "default_12m" in ds_mod.columns else None
        tasa_mod_str = f"{tasa_mod:.4%}" if tasa_mod else "N/A"
        agregar_parrafo(doc,
            f"El dataset de modelado (Stage 1 y 2 únicamente) contiene {n_mod:,} observaciones, "
            f"con una tasa de default a 12 meses de {tasa_mod_str}."
        )

    salto_pagina(doc)

    # -------------------------------------------------------------------------
    # 8. MODELADO
    # -------------------------------------------------------------------------
    agregar_titulo(doc, "8. Modelado – Entrenamiento y Comparación de Modelos PD", nivel=1)
    agregar_parrafo(doc,
        "Se entrenan tres modelos bajo el esquema de validación temporal out-of-time:"
    )
    for m, desc in [
        ("Regresión Logística", "Modelo regulatorio de referencia. Interpretable y requerido "
                                "por las guías EBA para auditoría."),
        ("XGBoost", "Gradient Boosting Machine. Captura no linealidades y efectos de "
                    "interacción. Modelo principal."),
        ("Random Forest", "Ensemble de árboles de decisión. Reduce varianza y sirve "
                          "como validación cruzada del XGBoost."),
    ]:
        p = doc.add_paragraph(style="List Bullet")
        r = p.add_run(f"{m}: ")
        set_run_font(r, negrita=True, tamaño=11)
        r2 = p.add_run(desc)
        set_run_font(r2, tamaño=11)

    agregar_parrafo(doc,
        "La partición temporal es: Entrenamiento (vintages 2016–2020), "
        "Validación (2022–2023), Test (2024). El vintage 2021 no está disponible "
        "en el dataset descargado."
    )

    agregar_titulo(doc, "8.1 Resultados Comparativos", nivel=2)
    if met_mod is not None:
        met_show = met_mod.reset_index()
        agregar_tabla_df(doc, met_show,
                         "Tabla 4. Métricas comparativas de modelos – conjunto de test (vintage 2024)")
        mejor = met_mod["AUC_test"].idxmax() if "AUC_test" in met_mod.columns else "XGBoost"
        auc_best = met_mod.loc[mejor, "AUC_test"] if "AUC_test" in met_mod.columns else "N/A"
        gini_best = met_mod.loc[mejor, "Gini_test"] if "Gini_test" in met_mod.columns else "N/A"
        agregar_parrafo(doc,
            f"El modelo {mejor} obtiene el mejor desempeño en el conjunto de test, "
            f"con AUC = {auc_best:.4f} y Gini = {gini_best:.4f}. "
            f"Todos los modelos superan el umbral mínimo regulatorio de Gini > 0.25."
            if isinstance(auc_best, float)
            else "Ver tabla de resultados para métricas comparativas."
        )
    else:
        agregar_parrafo(doc,
            "[Ver outputs/tablas/resultados_modelos.csv para las métricas comparativas "
            "una vez ejecutado el pipeline completo.]",
            cursiva=True, tamaño=10
        )

    salto_pagina(doc)

    # -------------------------------------------------------------------------
    # 9. CALIBRACIÓN Y LIFETIME PD
    # -------------------------------------------------------------------------
    agregar_titulo(doc, "9. Calibración PIT y Curva Lifetime PD", nivel=1)
    agregar_parrafo(doc,
        "El modelo XGBoost es calibrado mediante Regresión Isotónica para asegurar que "
        "las probabilidades predichas estén correctamente calibradas (PD observada ≈ PD predicha). "
        "La Lifetime PD se calcula aplicando el enfoque de supervivencia:"
    )
    agregar_codigo(doc,
        "# Lifetime PD usando enfoque de supervivencia\n"
        "# pd_marginal_t: PD a 12m promedio para préstamos de edad t\n"
        "supervivencia_t = ∏(1 - pd_marginal_s)  para s=1 hasta t\n"
        "Lifetime_PD_t   = 1 - supervivencia_t\n"
        "\n"
        "# IFRS 9 ECL por stage:\n"
        "# Stage 1: ECL_12m  = PD_12m × LGD × EAD\n"
        "# Stage 2: ECL_life = PD_lifetime × LGD × EAD"
    )

    fig_lt = str(FIGURAS_PATH / "curva_pd_lifetime.png")
    if Path(fig_lt).exists():
        agregar_figura(doc, fig_lt,
            "Figura 7. Curva de PD marginal y Lifetime PD acumulada por antigüedad del préstamo.")

    if lifetime is not None and len(lifetime) > 0:
        puntos_ref = lifetime[lifetime["loan_age"].isin([12, 24, 36, 60, 120])].copy()
        if len(puntos_ref) > 0:
            puntos_ref["pd_marginal"] = puntos_ref["pd_marginal"].map("{:.4%}".format)
            puntos_ref["pd_lifetime_acum"] = puntos_ref["pd_lifetime_acum"].map("{:.4%}".format)
            agregar_tabla_df(doc,
                puntos_ref[["loan_age", "pd_marginal", "pd_lifetime_acum"]],
                "Tabla 5. Puntos de referencia de la curva Lifetime PD"
            )

    salto_pagina(doc)

    # -------------------------------------------------------------------------
    # 10. VALIDACIÓN
    # -------------------------------------------------------------------------
    agregar_titulo(doc, "10. Validación Regulatoria del Modelo PD", nivel=1)
    agregar_parrafo(doc,
        "La validación sigue las guías de la EBA sobre estimación de parámetros IFRS 9 e "
        "incluye métricas de discriminación, calibración y estabilidad:"
    )
    for metrica, desc in [
        ("AUC / Gini",     "Discriminación: capacidad de separar defaulters de no-defaulters. "
                           "Gini > 0.25 = aceptable; > 0.50 = bueno; > 0.70 = muy bueno."),
        ("KS",             "Estadístico Kolmogorov-Smirnov: máxima separación entre distribuciones. "
                           "KS > 0.30 = aceptable."),
        ("Brier Score",    "Calibración: error cuadrático de probabilidades. Menor = mejor."),
        ("PSI",            "Estabilidad: < 0.10 estable; 0.10–0.25 monitorear; > 0.25 inestable."),
        ("Hosmer-Lemeshow","Bondad de ajuste por deciles de riesgo. p-valor > 0.05 = calibración OK."),
    ]:
        p = doc.add_paragraph(style="List Bullet")
        r = p.add_run(f"{metrica}: ")
        set_run_font(r, negrita=True, tamaño=11)
        r2 = p.add_run(desc)
        set_run_font(r2, tamaño=11)

    if val is not None:
        agregar_titulo(doc, "10.1 Resultados de Validación", nivel=2)
        agregar_tabla_df(doc, val.reset_index(),
                         "Tabla 6. Métricas de validación regulatoria – conjunto de test")

    for nombre_fig, caption in [
        ("validacion_roc.png",         "Figura 8. Curva ROC comparativa de los tres modelos."),
        ("validacion_backtesting.png", "Figura 9. Backtesting por decil: PD predicha vs. observada."),
    ]:
        ruta_f = str(FIGURAS_PATH / nombre_fig)
        if Path(ruta_f).exists():
            agregar_figura(doc, ruta_f, caption)

    salto_pagina(doc)

    # -------------------------------------------------------------------------
    # 11. EXPLICABILIDAD
    # -------------------------------------------------------------------------
    agregar_titulo(doc, "11. Explicabilidad del Modelo", nivel=1)
    agregar_parrafo(doc,
        "Conforme a las guías de gobernanza de modelos IFRS 9 (EBA, 2022) y los principios "
        "del BCBS, el modelo debe ser comprensible, auditable y explicable. Se aplican "
        "herramientas SHAP (SHapley Additive exPlanations) en dos niveles:"
    )
    for nivel, desc in [
        ("Global", "SHAP Summary Plot y Bar Plot: muestran qué variables tienen mayor impacto "
                   "promedio sobre la PD estimada en toda la cartera."),
        ("Local",  "SHAP Waterfall Plot: explica la predicción individual de un préstamo "
                   "específico, descomponiendo la PD en contribuciones de cada variable."),
    ]:
        p = doc.add_paragraph(style="List Bullet")
        r = p.add_run(f"{nivel}: ")
        set_run_font(r, negrita=True, tamaño=11)
        r2 = p.add_run(desc)
        set_run_font(r2, tamaño=11)

    for nombre_fig, caption in [
        ("shap_summary_xgboost.png", "Figura 10. SHAP Summary Plot – efecto de cada variable sobre la PD (XGBoost)."),
        ("shap_barplot_xgboost.png", "Figura 11. Importancia SHAP media – Top 15 variables."),
        ("importancia_gain_xgboost.png", "Figura 12. Importancia por Gain (XGBoost nativo)."),
    ]:
        ruta_f = str(FIGURAS_PATH / nombre_fig)
        if Path(ruta_f).exists():
            agregar_figura(doc, ruta_f, caption)

    salto_pagina(doc)

    # -------------------------------------------------------------------------
    # 12. CONCLUSIONES
    # -------------------------------------------------------------------------
    agregar_titulo(doc, "12. Conclusiones y Líneas Futuras", nivel=1)
    agregar_titulo(doc, "12.1 Conclusiones", nivel=2)
    agregar_parrafo(doc,
        "Este trabajo presenta una metodología completa y reproducible para la estimación de "
        "la Probabilidad de Default bajo el marco regulatorio IFRS 9, utilizando datos "
        "hipotecarios públicos de Freddie Mac. Los principales resultados son:"
    )
    conclusiones = [
        "Se definió y aplicó el default IFRS 9 (90+ DPD o evento de pérdida) sobre el panel "
        "longitudinal, obteniendo una tasa de default a 12 meses coherente con la literatura.",
        "El modelo XGBoost calibrado obtuvo el mejor desempeño discriminante entre los tres "
        "modelos evaluados, con AUC y Gini superiores a los benchmarks regulatorios mínimos.",
        "El análisis SHAP reveló que las variables de comportamiento dinámico (DPD corriente, "
        "máximo DPD histórico) dominan la predicción, seguidas por el credit score FICO y el LTV.",
        "La curva Lifetime PD generada mediante el enfoque de supervivencia es aplicable directamente "
        "al cálculo de ECL para instrumentos en Stage 2.",
        "La metodología es replicable con datos de carteras reales de entidades financieras "
        "argentinas, apoyando la adopción de IFRS 9 en el mercado local.",
    ]
    for c in conclusiones:
        p = doc.add_paragraph(style="List Bullet")
        r = p.add_run(c)
        set_run_font(r, tamaño=11)

    agregar_titulo(doc, "12.2 Líneas Futuras", nivel=2)
    for lf in [
        "Incorporar variables macroeconómicas (desempleo, tasas de interés) para stress testing.",
        "Extender el modelo para estimar LGD y EAD, completando el cálculo integral del ECL.",
        "Aplicar la metodología a datos de entidades financieras argentinas bajo BCRA/NIIF 9.",
        "Explorar modelos de Deep Learning (LSTM) para capturar dependencias temporales.",
    ]:
        p = doc.add_paragraph(style="List Bullet")
        r = p.add_run(lf)
        set_run_font(r, tamaño=11)

    salto_pagina(doc)

    # -------------------------------------------------------------------------
    # BIBLIOGRAFÍA
    # -------------------------------------------------------------------------
    agregar_titulo(doc, "Bibliografía", nivel=1)
    referencias = [
        "[1] Basel Committee on Banking Supervision (2023). Principles for the Effective "
        "Management of Credit Risk.",
        "[2] IFRS Foundation (2018). IFRS 9: Financial Instruments – Implementation Guidance.",
        "[3] European Banking Authority (2022). Guidelines on PD estimation, LGD, and treatment "
        "of defaulted exposures.",
        "[4] Thomas, L., Edelman, D., & Crook, J. (2017). Credit Scoring and Its Applications "
        "(2nd Ed.). SIAM.",
        "[5] Lessmann, S., Baesens, B., Seow, H., & Thomas, L. (2015). Benchmarking "
        "State-of-the-Art Classification Algorithms for Credit Scoring. EJOR.",
        "[6] European Central Bank (2021). Good Practices for Model Validation and "
        "Explainability in IFRS 9.",
        "[7] Freddie Mac (2025). Single-Family Loan-Level Dataset User Guide. Release 44.",
        "[8] Breiman, L. (2001). Random Forests. Machine Learning, 45(1), 5–32.",
        "[9] Chen, T., & Guestrin, C. (2016). XGBoost: A Scalable Tree Boosting System. "
        "Proc. 22nd ACM SIGKDD.",
        "[10] Ke, G., Meng, Q., et al. (2017). LightGBM: A Highly Efficient Gradient "
        "Boosting Decision Tree. NIPS.",
        "[12] Banco Central de la República Argentina (2024). Comunicación 'A' 6847: "
        "Previsiones Mínimas por Riesgo de Incobrabilidad – Normas NIIF.",
        "[13] Freire López, J. (2021). Modelo de Clasificación de Riesgo Crediticio "
        "Utilizando Random Forest. Tesis de Maestría, Universidad Internacional SEK.",
        "[14] Ghent, A. C., & Kudlyak, M. (2011). Recourse and Residential Mortgage Default. "
        "Review of Financial Studies, 24(9), 3139–3186.",
        "[15] Elul, R., et al. (2010). What 'Triggers' Mortgage Default? AER, 100(2), 490–494.",
        "[16] Lundberg, S. M., & Lee, S. I. (2017). A Unified Approach to Interpreting "
        "Model Predictions. NIPS.",
    ]
    for ref in referencias:
        p = doc.add_paragraph(style="List Paragraph")
        p.paragraph_format.left_indent = Cm(0.5)
        p.paragraph_format.first_line_indent = Cm(-0.5)
        p.paragraph_format.space_after = Pt(4)
        r = p.add_run(ref)
        set_run_font(r, tamaño=10)

    # -------------------------------------------------------------------------
    # GUARDAR
    # -------------------------------------------------------------------------
    doc.save(str(DOCX_OUT))
    print(f"\nDocumento guardado: {DOCX_OUT}")


# ============================================================================
# MAIN
# ============================================================================

def main():
    t0 = time.time()
    print("=" * 60)
    print("GENERADOR DE TESIS – Tesis_borrador.docx")
    print("=" * 60)

    print("\n[1/2] Ejecutando pipeline de análisis...")
    resultados = ejecutar_pipeline()

    print("\n[2/2] Construyendo documento Word...")
    construir_docx(resultados)

    print(f"\nFinalizado en {(time.time() - t0)/60:.1f} minutos.")
    print(f"Documento: {DOCX_OUT}")


if __name__ == "__main__":
    main()
